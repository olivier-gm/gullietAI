import subprocess
import os
import string
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class Document_process:
    
    @staticmethod
    def convert(input_file, output_folder):
        LIBRE = 'C:/Program Files/LibreOffice/program/soffice.exe'
        commandStrings = [LIBRE, "--headless", "--convert-to",
                           "pdf", "--outdir", output_folder, input_file]
        retCode = subprocess.call(commandStrings)

        if retCode == 0:
            print('conversion successful!')
        else:
            print(f'conversion error. Code: {retCode}')

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def llenar_campos(replacements, document):
        for paragraph in document.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

    @staticmethod
    def fecha(date):
        string = ''
        if date != '':
            day = date[0:2]
            month = date[3:5]
            year = date[6:10]
            months = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio',
                    'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
            new_month = months[int(month) - 1]
            string = f'{day} de {new_month} de {year}'

        return string

    @staticmethod
    def capitalizar_frases(cadena):
        palabras = cadena.split()
        palabras_capitalizadas = []
        for palabra in palabras:
            if palabra.lower() in ['de', 'del', 'la', 'las', 'los', 'y', 'para'] and palabras_capitalizadas != []:
                palabras_capitalizadas.append(palabra.lower())
            else:
                palabras_capitalizadas.append(palabra.capitalize())

        nueva_cadena = ' '.join(palabras_capitalizadas)
        return nueva_cadena

    @staticmethod
    def generate_random_code(length=2):
        letters_and_digits = string.ascii_letters + string.digits
        return ''.join(random.choice(letters_and_digits) for i in range(length))

    @staticmethod
    def docx_replace(doc, old_text, new_text):
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    Document_process.docx_replace(cell, old_text, new_text)

    @staticmethod
    def remove_file(file_path):
        base_path = os.path.splitext(file_path)[0]
        for extension in ['.pdf', '.docx']:
            try:
                os.remove(base_path + extension)
            except Exception as error:
                print("Error removing or closing downloaded file handle", error)

    @staticmethod
    def parrafos(body, document, topic, flag):
        if body != '':
            p = document.add_paragraph(topic)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.paragraph_format.line_spacing = Pt(21)
            document.add_paragraph('').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            o = 0
            body = body.split('\n')
            for i in body:
                if i == '':
                    body.pop(o)
                o += 1
            for parrafo in body:
                p = document.add_paragraph(parrafo)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.paragraph_format.line_spacing = Pt(21)
            Document_process.docx_replace(document, '\r', '')
            Document_process.docx_replace(document, '\n', '')
            if flag == True:
                document.add_page_break()

    @staticmethod
    def underline_words_in_first_page(doc, words):
        for i in range(24, 37):
            paragraph = doc.paragraphs[i]
            new_runs = []
            for run in paragraph.runs:
                found = False
                for word in words:
                    if word in run.text:
                        parts = run.text.split(word)
                        for part in parts[:-1]:
                            new_runs.append((part, run.bold, run.italic, run.underline))
                            new_runs.append((word, run.bold, run.italic, True))
                        new_runs.append((parts[-1], run.bold, run.italic, run.underline))
                        found = True
                        break
                if not found:
                    new_runs.append((run.text, run.bold, run.italic, run.underline))
            for run in paragraph.runs:
                run.clear()
            for text, bold, italic, underline in new_runs:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
                run.underline = underline

    @staticmethod
    def fill_placeholders(docx_output, template_path, template_path2, replacements,
                           introduction, essay_content, conclusion, head_title, id):
        if id == 'bach':
            words = ['DOCENTE:', 'ALUMNOS:','ALUMNO:', 'MATERIA:']
        else:
            words = ['DOCENTE:', 'ALUMNOS:','ALUMNO:', 'SECCION:', 'AÑO:', 'SEMESTRE:', 'TRIMESTRE:', 'MATERIA:']
            
        if essay_content == '' and introduction == '' and conclusion == '':
            document = Document(template_path2)
        else:
            document = Document(template_path)

        Document_process.llenar_campos(replacements, document)
        Document_process.underline_words_in_first_page(document, words)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        paragraph = document.paragraphs[17]
        run = paragraph.runs[0]
        run.font.size = Pt(17.5)

        if essay_content != '' or introduction != '' or conclusion != '':
            document.add_paragraph('')
            document.add_paragraph('')

        flagi = False
        if essay_content != '' or conclusion != '':
            flagi = True

        flage = False
        if conclusion != '':
            flage = True

        Document_process.parrafos(introduction, document, 'Introducción', flagi)
        Document_process.parrafos(essay_content, document, head_title, flage)
        Document_process.parrafos(conclusion, document, 'Conclusión', False)

        head_title = head_title.replace(':', '_')
        document.save(docx_output)
        Document_process.convert(docx_output, 'output')